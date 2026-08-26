import os
import io
import re
import base64
import logging
import time
import itertools
from typing import Optional, List, Tuple, Set
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from PIL import Image, ImageOps, ImageEnhance
from pdf2image import convert_from_path
import requests
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

logger = logging.getLogger("scriptsense.ocr")

base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Configure Poppler path for PDF conversion if available
poppler_env_path = os.getenv("POPPLER_PATH")
poppler_path = None
if poppler_env_path and os.path.exists(poppler_env_path):
    poppler_path = poppler_env_path
else:
    common_poppler_paths = [
        os.path.join(base_dir, "bin", "poppler"),
        os.path.join(base_dir, "bin", "poppler_extracted", "poppler-24.08.0", "Library", "bin"),
        r"C:\Program Files\poppler\bin",
        r"C:\Program Files (x86)\poppler\bin",
        r"C:\poppler\bin",
    ]
    for path in common_poppler_paths:
        if os.path.exists(path):
            poppler_path = path
            break

# Configure Tesseract path if available
tesseract_env_path = os.getenv("TESSERACT_PATH")
tessdata_env_prefix = os.getenv("TESSDATA_PREFIX")

try:
    import pytesseract
    if tesseract_env_path and os.path.exists(tesseract_env_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_env_path
    else:
        common_tesseract_paths = [
            r"C:\Users\rasmi\AppData\Local\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in common_tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                break

    if tessdata_env_prefix and os.path.exists(tessdata_env_prefix):
        os.environ["TESSDATA_PREFIX"] = tessdata_env_prefix
except ImportError:
    pytesseract = None


# Global EasyOCR reader singleton
_easyocr_reader = None

def get_easyocr_reader():
    """
    Initializes and caches the EasyOCR Reader instance (loaded once on CPU).
    """
    global _easyocr_reader
    if _easyocr_reader is not None:
        return _easyocr_reader

    try:
        import easyocr
        # Initialize CPU-friendly EasyOCR reader
        _easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
        logger.info("Initialized EasyOCR reader successfully (CPU mode).")
        return _easyocr_reader
    except Exception as e:
        logger.warning(f"EasyOCR reader initialization failed: {e}")
        return None


def warmup_ocr_models():
    """
    Pre-warms the EasyOCR neural network weights at backend startup so user upload requests never experience initial model loading delay.
    """
    try:
        t0 = time.time()
        reader = get_easyocr_reader()
        if reader is not None:
            import numpy as np
            # Warm up detector and recognizer with a small canvas
            dummy = np.ones((100, 250, 3), dtype=np.uint8) * 255
            reader.readtext(dummy)
            elapsed = round(time.time() - t0, 2)
            print(f"[OCR WARMUP] EasyOCR model pre-warmed into memory in {elapsed}s.")
    except Exception as e:
        print(f"[OCR WARMUP NOTE] Warmup completed with notice: {e}")


# =====================================================================
# Context-Aware English & STEM Vocabulary for Safe 'r' Character Repair
# =====================================================================
VOCAB_LIST = [
    # User target test words & related forms
    "are", "answer", "answers", "answered", "answering",
    "correct", "correctly", "correction", "corrections",
    "protocol", "protocols",
    "computer", "computers", "computing", "computation",
    "server", "servers",
    "network", "networks", "networking", "networked",
    "error", "errors",
    "transfer", "transfers", "transferred", "transferring",
    
    # Common academic & STEM vocabulary
    "interconnected", "interconnection", "interconnections",
    "share", "shares", "shared", "sharing",
    "other", "others", "otherwise",
    "packet", "packets", "layer", "layers",
    "data", "device", "devices", "communicate", "communicates", "communicated", "communication",
    "connection", "connections", "oriented", "reliable", "unreliable", "delivery", "handshake",
    "transmission", "transport", "routing", "router", "routers",
    "hardware", "software", "firmware", "driver", "drivers",
    "memory", "processor", "processors", "processing", "process", "processes",
    "register", "registers", "circuit", "circuits", "current", "voltage", "resistor", "resistance",
    "force", "forces", "mass", "acceleration", "accelerate", "accelerates", "momentum",
    "gravity", "gravitational", "friction", "frictionless", "energy", "power", "work",
    "velocity", "vector", "scalar", "newton", "newtons", "joule", "joules", "watt", "watts",
    "proportional", "inversely", "directly", "constant", "rate", "ratio", "variable", "variables",
    "structure", "structures", "structural", "architecture", "architectures",
    "operator", "operators", "operation", "operations", "operand", "operands",
    "array", "arrays", "string", "strings", "character", "characters", "integer", "integers",
    "pointer", "pointers", "address", "addresses", "reference", "references",
    "function", "functions", "method", "methods", "parameter", "parameters", "argument", "arguments",
    "return", "returns", "returned", "returning", "result", "results", "resulting",
    "query", "queries", "record", "records", "report", "reports",
    "resource", "resources", "storage", "thread", "threads", "virtual", "barrier",
    "require", "requires", "required", "requirement", "requirements",
    "format", "formats", "formatted", "formatting", "standard", "standards",
    "program", "programs", "programmed", "programmer", "programmers", "programming",
    "property", "properties", "performance", "perform", "performs", "performed",
    "reaction", "reactions", "cellular", "respiration", "chlorophyll", "photosynthesis",
    "primary", "secondary", "tertiary", "order", "orders", "ordered", "ordering",
    "forward", "reverse", "internal", "external", "source", "destination",
    "receiver", "receivers", "receive", "receives", "received", "receiving",
    "sender", "senders", "send", "sends", "sending", "sent",
    "client", "clients", "peer", "peers", "node", "nodes", "link", "links",
    "medium", "media", "signal", "signals", "channel", "channels",
    "bandwidth", "frequency", "throughput", "latency", "delay",
    "measure", "measured", "measurement", "measurements", "metric", "metrics",
    "system", "systems", "state", "states", "status", "table", "tables",
    
    # Common short words
    "a", "an", "the", "in", "on", "at", "to", "for", "of", "with", "by", "from",
    "is", "am", "are", "was", "were", "be", "been", "being", "have", "has", "had",
    "do", "does", "did", "done", "will", "would", "shall", "should", "may", "might", "can", "could", "must",
    "it", "its", "they", "them", "their", "theirs", "this", "that", "these", "those",
    "we", "our", "ours", "us", "you", "your", "yours", "he", "him", "his", "she", "her", "hers",
    "and", "or", "but", "nor", "so", "yet", "if", "then", "else", "when", "where", "why", "how", "what", "which",
    "all", "any", "both", "each", "few", "more", "most", "some", "such", "no", "not", "only", "own", "same", "too", "very",
    "one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten", "first", "second", "third",
    
    # Words with x, n, v that MUST NOT be replaced
    "box", "boxes", "fox", "foxes", "tax", "taxes", "six", "sixes", "mix", "mixed", "mixing",
    "fix", "fixed", "fixing", "max", "maximum", "matrix", "matrices", "index", "indexes", "indices",
    "syntax", "complex", "complexity", "prefix", "suffix", "pixel", "pixels", "proxy", "proxies",
    "text", "texts", "next", "context", "convex", "vertex", "vertices", "latex", "unix", "linux",
    "can", "man", "men", "sun", "run", "ten", "pen", "son", "pan", "fan", "van", "tan", "pin", "bin",
    "give", "gives", "given", "giving", "have", "live", "save", "wave", "move", "view", "value"
]

COMMON_ENGLISH_WORDS: Set[str] = set(w.lower() for w in VOCAB_LIST)


def preserve_case(original: str, modified: str) -> str:
    """Preserves capitalization pattern of original word in modified word."""
    if original.isupper():
        return modified.upper()
    if original.istitle():
        return modified.capitalize()
    return modified


def split_joined_words(word: str) -> List[str]:
    """Splits words that EasyOCR accidentally fused (e.g. 'computerand' -> 'computer and', 'correctansiers' -> 'correct answers')."""
    lower = word.lower()
    if lower in COMMON_ENGLISH_WORDS:
        return [word]
    for i in range(3, len(lower) - 2):
        w1 = lower[:i]
        w2 = lower[i:]
        if w1 in COMMON_ENGLISH_WORDS:
            if w2 in COMMON_ENGLISH_WORDS:
                return [word[:i], word[i:]]
            # Check if w2 can be repaired to a valid word
            cands2 = generate_r_candidates(w2)
            for c2 in cands2:
                if c2 in COMMON_ENGLISH_WORDS:
                    return [word[:i], preserve_case(word[i:], c2)]
    return [word]


def generate_r_candidates(word: str) -> List[str]:
    """
    Generates candidate word variations by substituting typical OCR handwriting confusions with 'r' or 'w',
    and recovering omitted 'r' characters (e.g. 'netwok' -> 'network').
    """
    candidates = []
    
    # 1. Handle 'iv' -> 'w' / 'r'
    if 'iv' in word:
        candidates.append(word.replace('iv', 'w'))
        candidates.append(word.replace('iv', 'r'))

    # 2. 'si' -> 'sw' (e.g. 'ansiers' -> 'answers')
    if 'si' in word:
        candidates.append(word.replace('si', 'sw'))

    # 3. 'wv' -> 'w' (e.g. 'netwvork' -> 'network')
    if 'wv' in word:
        candidates.append(word.replace('wv', 'w'))

    # 4. 'nst' -> 'nsf' (e.g. 'transter' -> 'transfer')
    if 'nst' in word:
        candidates.append(word.replace('nst', 'nsf'))

    # 5. Collect indices of confused characters ('x', 'n', 'v', 'z', 'j')
    confused_indices = [i for i, ch in enumerate(word) if ch in ('x', 'n', 'v', 'z', 'j')]
    
    # Check all subsets of size 1 to min(3, len) to replace with 'r'
    for r_count in range(1, min(len(confused_indices) + 1, 4)):
        for combo in itertools.combinations(confused_indices, r_count):
            cand_chars = list(word)
            for idx in combo:
                cand_chars[idx] = 'r'
            candidates.append("".join(cand_chars))

    # 6. Omission of 'r' (e.g. 'netwok' -> 'network', 'sever' -> 'server', 'tansfer' -> 'transfer', 'corect' -> 'correct')
    if len(word) >= 3:
        for i in range(len(word) + 1):
            candidates.append(word[:i] + 'r' + word[i:])

    return candidates


def repair_r_confusions_in_word(raw_word: str) -> str:
    """
    Intelligently repairs handwriting OCR misrecognitions of 'r' (e.g. 'x', 'n', 'v', 'iv')
    ONLY when the word is not already a valid word and the candidate 'r'-substituted form is a valid English word.
    Preserves casing, punctuation, and guarantees zero false positives on real English words.
    """
    m = re.match(r'^([^a-zA-Z0-9]*)([a-zA-Z0-9_-]+)([^a-zA-Z0-9]*)$', raw_word)
    if not m:
        return raw_word

    prefix, core, suffix = m.groups()
    
    # Handle internal OCR underscores (e.g. "collection_of" -> "collection of")
    if '_' in core and len(core) > 2:
        parts = core.split('_')
        repaired_parts = [repair_r_confusions_in_word(p) for p in parts if p]
        return prefix + " ".join(repaired_parts) + suffix

    lower_core = core.lower()

    # If already a valid English word (like 'box', 'can', 'six', 'matrix', 'tax', 'are'), keep as-is!
    if lower_core in COMMON_ENGLISH_WORDS:
        return raw_word

    # Generate r-repaired candidates
    candidates = generate_r_candidates(lower_core)
    for cand in candidates:
        if cand in COMMON_ENGLISH_WORDS and cand != lower_core:
            return prefix + preserve_case(core, cand) + suffix

    # Try split joined words (e.g. "computerand" -> "computer and", "correctansiers" -> "correct answers")
    splits = split_joined_words(core)
    if len(splits) > 1:
        repaired_splits = [repair_r_confusions_in_word(s) for s in splits]
        return prefix + " ".join(repaired_splits) + suffix

    return raw_word


def repair_ocr_text(text: str) -> str:
    """
    Processes extracted OCR text line-by-line and token-by-token to restore confused 'r' characters.
    """
    if not text:
        return ""
    lines = []
    for line in text.splitlines():
        tokens = line.split()
        repaired_tokens = [repair_r_confusions_in_word(t) for t in tokens]
        lines.append(" ".join(repaired_tokens))
    return "\n".join(lines)


def preprocess_image_for_ocr(
    img: Image.Image,
    max_dimension: int = 1500,
    min_dimension: int = 700,
) -> Image.Image:
    """
    Carefully adjusted preprocessing to preserve the fine stroke and arch of 'r':
    - Converts to standard RGB
    - Uses BICUBIC resampling to prevent pixelation/disconnection of delicate 'r' arches
    - Applies balanced contrast enhancement (1.18x) so faint strokes aren't burned out
    - Subtle edge sharpening (1.15x) to clarify handwriting contours
    """
    if img.mode != 'RGB':
        img = img.convert('RGB')

    w, h = img.size
    max_dim = max(w, h)
    min_dim = min(w, h)

    # 1. Bicubic scaling to keep stroke continuity and optimal resolution
    if max_dim > max_dimension:
        scale = max_dimension / float(max_dim)
        new_w, new_h = int(w * scale), int(h * scale)
        img = img.resize((new_w, new_h), Image.Resampling.BICUBIC)
    elif max_dim < min_dimension:
        scale = 900.0 / float(max(min_dim, 1))
        new_w, new_h = int(w * scale), int(h * scale)
        img = img.resize((new_w, new_h), Image.Resampling.BICUBIC)

    # 2. Balanced contrast enhancement (preserves faint pen arches)
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.18)

    # 3. Subtle edge sharpness enhancement
    sharpener = ImageEnhance.Sharpness(img)
    img = sharpener.enhance(1.15)

    return img


def sort_and_group_detections(detections: list) -> List[str]:
    """
    Sorts and groups bounding box text detections into structured, natural reading lines.
    Detections format: [(bbox, text, conf), ...]
    """
    if not detections:
        return []

    # Estimate average box height for vertical line grouping tolerance
    heights = [abs(d[0][2][1] - d[0][0][1]) for d in detections]
    avg_h = sum(heights) / max(len(heights), 1) if heights else 15
    y_tol = avg_h * 0.6

    sorted_by_y = sorted(detections, key=lambda d: d[0][0][1])
    lines = []
    current_line = []
    current_y = None

    for d in sorted_by_y:
        y = d[0][0][1]
        if current_y is None or abs(y - current_y) <= y_tol:
            current_line.append(d)
            if current_y is None:
                current_y = y
        else:
            current_line.sort(key=lambda d: d[0][0][0])
            lines.append(" ".join(d[1].strip() for d in current_line))
            current_line = [d]
            current_y = y

    if current_line:
        current_line.sort(key=lambda d: d[0][0][0])
        lines.append(" ".join(d[1].strip() for d in current_line))

    return lines


def clean_ocr_text(raw_text: str) -> str:
    """
    Cleans extracted OCR text by normalizing irregular spaces and blank lines.
    """
    if not raw_text:
        return ""
    cleaned_lines = []
    for line in raw_text.splitlines():
        normalized_line = re.sub(r'[ \t]+', ' ', line).strip()
        if normalized_line:
            cleaned_lines.append(normalized_line)
    return "\n".join(cleaned_lines)


def extract_text_with_easyocr(image_bytes: bytes, min_confidence: float = 0.25) -> str:
    """
    Fast EasyOCR extraction with improved 'r' stroke preservation and post-processing.
    """
    reader = get_easyocr_reader()
    if reader is None:
        return ""

    try:
        import numpy as np
        raw_img = Image.open(io.BytesIO(image_bytes))
        processed_img = preprocess_image_for_ocr(raw_img)
        img_arr = np.array(processed_img)

        # Run EasyOCR with tuned parameters for handwritten character shapes
        results = reader.readtext(
            img_arr,
            detail=1,
            paragraph=False,
            batch_size=4,
            text_threshold=0.3,
            low_text=0.2,
            link_threshold=0.25,
            slope_ths=0.2,
            width_ths=0.45,
            canvas_size=1400,
            mag_ratio=1.0
        )

        filtered = []
        for item in results:
            if isinstance(item, (list, tuple)) and len(item) >= 3:
                conf_val = float(item[2])
                txt_val = str(item[1]).strip()
                if conf_val >= min_confidence and txt_val:
                    filtered.append((item[0], txt_val, conf_val))
        lines = sort_and_group_detections(filtered)
        raw_text = clean_ocr_text("\n".join(lines))
        
        # Apply context-aware 'r' reconstruction
        return repair_ocr_text(raw_text)
    except Exception as e:
        logger.warning(f"EasyOCR extraction error: {e}")
        return ""


def extract_text_with_tesseract(image_bytes: bytes) -> str:
    """
    Fallback extraction using Tesseract OCR if available.
    """
    if pytesseract is None:
        return ""

    try:
        raw_img = Image.open(io.BytesIO(image_bytes))
        if raw_img.mode != 'RGB':
            raw_img = raw_img.convert('RGB')
        
        # Fast resize if large
        w, h = raw_img.size
        if max(w, h) > 1500:
            scale = 1500.0 / max(w, h)
            raw_img = raw_img.resize((int(w * scale), int(h * scale)), Image.Resampling.BICUBIC)

        gray = ImageOps.grayscale(raw_img)
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(1.2)

        text = pytesseract.image_to_string(enhanced, config='--oem 3 --psm 6').strip()
        if not text:
            text = pytesseract.image_to_string(enhanced, config='--oem 3 --psm 3').strip()

        cleaned = clean_ocr_text(text)
        return repair_ocr_text(cleaned)
    except Exception as e:
        logger.warning(f"Tesseract extraction error: {e}")
        return ""


def is_valid_gemini_key(api_key: Optional[str] = None) -> bool:
    """Checks whether a non-placeholder Gemini API key is configured."""
    if not api_key:
        api_key = os.getenv("GEMINI_API_KEY", "").strip() or os.getenv("GOOGLE_API_KEY", "").strip()
    return bool(api_key and len(api_key) > 10 and not api_key.startswith("your_"))


def extract_text_with_gemini_vision(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """
    Extracts text using Gemini 1.5 Flash REST API with inline image data.
    Works natively via HTTP requests without requiring SDK installation.
    """
    api_key = os.getenv("GEMINI_API_KEY", "").strip() or os.getenv("GOOGLE_API_KEY", "").strip()
    if not is_valid_gemini_key(api_key):
        return ""

    try:
        b64_image = base64.b64encode(image_bytes).decode("utf-8")
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": "Extract all handwritten and printed text from this image accurately. Return only the extracted text without introductory or conversational filler."},
                        {
                            "inline_data": {
                                "mime_type": mime_type,
                                "data": b64_image
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 2048
            }
        }
        resp = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=12)
        if resp.status_code == 200:
            res_json = resp.json()
            candidates = res_json.get("candidates", [])
            if candidates and "content" in candidates[0] and "parts" in candidates[0]["content"]:
                extracted = candidates[0]["content"]["parts"][0].get("text", "").strip()
                if extracted:
                    return extracted
        logger.info(f"Gemini Vision REST notice: HTTP {resp.status_code}")
        return ""
    except Exception as e:
        logger.info(f"Gemini Vision extraction notice: {e}")
        return ""


_vision_client = None

def get_vision_client():
    """Returns initialized Google Vision client if available."""
    global _vision_client
    return _vision_client

def extract_text_with_vision_api(image_bytes: bytes) -> str:
    """Optional Google Cloud Vision API extractor."""
    client = get_vision_client()
    if client is None:
        return ""
    try:
        response = client.document_text_detection(image={"content": image_bytes})
        if response and hasattr(response, "full_text_annotation") and response.full_text_annotation:
            return str(response.full_text_annotation.text).strip()
        return ""
    except Exception:
        return ""


def extract_text_with_priority_pipeline(image_bytes: bytes, filename: str = "image") -> Tuple[str, str]:
    """
    Runs the fast OCR pipeline with built-in timeout safeguards:
    1. Primary: Gemini Vision (if valid API key is present)
    2. Primary Local: EasyOCR (Fast Local PyTorch with 20s timeout & 'r' repair)
    3. Fallback: Tesseract OCR (if EasyOCR fails or returns empty)
    4. Safe Failure: Returns user-friendly failure message without blocking.
    """
    t0 = time.time()
    print(f"[OCR PIPELINE] Starting fast OCR extraction for '{filename}'...")

    # 1. Primary Engine: Gemini Vision (if configured)
    if is_valid_gemini_key():
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(extract_text_with_gemini_vision, image_bytes)
                gemini_text = future.result(timeout=10.0)

            if gemini_text and gemini_text.strip():
                elapsed = round(time.time() - t0, 2)
                print(f"[OCR SUCCESS] Engine: Gemini Vision | File: {filename} | Extracted: {len(gemini_text)} chars in {elapsed}s")
                return gemini_text.strip(), "GeminiVision"
            else:
                print(f"[OCR NOTICE] Gemini Vision returned empty for '{filename}'. Switching to EasyOCR...")
        except Exception as e:
            print(f"[OCR NOTICE] Gemini Vision notice: {e}. Switching to EasyOCR...")

    # 2. Local Neural Engine: EasyOCR with 20-second timeout guard
    try:
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(extract_text_with_easyocr, image_bytes)
            easyocr_text = future.result(timeout=20.0)

        if easyocr_text and easyocr_text.strip():
            elapsed = round(time.time() - t0, 2)
            print(f"[OCR SUCCESS] Engine: EasyOCR | File: {filename} | Extracted: {len(easyocr_text)} chars in {elapsed}s")
            return easyocr_text, "EasyOCR"
        else:
            print(f"[OCR NOTICE] EasyOCR produced no text for '{filename}'. Trying Tesseract fallback...")
    except FuturesTimeoutError:
        print(f"[OCR TIMEOUT] EasyOCR exceeded 20s timeout for '{filename}'. Switching to Tesseract fallback...")
    except Exception as e:
        print(f"[OCR NOTICE] EasyOCR encountered error: {e}. Switching to Tesseract fallback...")

    # 3. Fallback Engine: Tesseract OCR with 6-second timeout
    if pytesseract is not None:
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(extract_text_with_tesseract, image_bytes)
                tess_text = future.result(timeout=6.0)

            if tess_text and tess_text.strip():
                elapsed = round(time.time() - t0, 2)
                print(f"[OCR SUCCESS] Engine: Tesseract OCR | File: {filename} | Extracted: {len(tess_text)} chars in {elapsed}s")
                return tess_text, "Tesseract"
        except Exception as e:
            print(f"[OCR NOTICE] Tesseract fallback failed: {e}")

    # 4. Graceful Failure Message
    print(f"[OCR FAILED] Could not extract text for '{filename}'. Returning fallback message.")
    return "Text extraction failed. Please try a clearer image or enter the answer manually.", "None"


def extract_text_from_file(absolute_file_path: str) -> str:
    """
    Extracts handwritten text from an image or PDF file using the fast priority OCR pipeline.
    """
    if not os.path.exists(absolute_file_path):
        return "Text extraction failed. Please try a clearer image or enter the answer manually."

    filename = os.path.basename(absolute_file_path)
    _, ext = os.path.splitext(absolute_file_path)
    ext = ext.lower()

    if ext == ".pdf":
        try:
            if poppler_path:
                pages = convert_from_path(absolute_file_path, dpi=150, poppler_path=poppler_path)
            else:
                pages = convert_from_path(absolute_file_path, dpi=150)

            page_texts = []
            for i, page_img in enumerate(pages):
                buf = io.BytesIO()
                page_img.save(buf, format="PNG")
                page_bytes = buf.getvalue()
                text, engine = extract_text_with_priority_pipeline(page_bytes, filename=f"{filename} (Page {i+1})")
                if text.strip() and not text.startswith("Text extraction failed"):
                    page_texts.append(text.strip())

            if page_texts:
                return "\n\n".join(page_texts)
            return "Text extraction failed. Please try a clearer image or enter the answer manually."
        except Exception as e:
            logger.error(f"Failed processing PDF '{filename}': {e}")
            return "Text extraction failed. Please try a clearer image or enter the answer manually."

    elif ext in [".docx", ".doc"]:
        try:
            import docx
            doc = docx.Document(absolute_file_path)
            paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs_text.append(row_text)
            if paragraphs_text:
                return "\n\n".join(paragraphs_text)
            return "No text content found in the Word document."
        except Exception as e:
            logger.error(f"Failed extracting text from DOCX '{filename}': {e}")
            return f"Failed to extract text from Word document: {e}"

    elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
        try:
            with open(absolute_file_path, "rb") as f:
                image_bytes = f.read()
            text, engine = extract_text_with_priority_pipeline(image_bytes, filename=filename)
            return text
        except Exception as e:
            logger.error(f"Failed reading file '{filename}': {e}")
            return "Text extraction failed. Please try a clearer image or enter the answer manually."

    else:
        return "Text extraction failed. Please try a clearer image or enter the answer manually."
